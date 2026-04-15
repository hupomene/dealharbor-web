from __future__ import annotations

from io import BytesIO
from typing import Any, Iterable, Sequence
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from .pdf_converter import convert_docx_bytes_to_pdf_bytes
from .schemas import DocumentPreviewPayload, OutputFormat
from .template_registry import get_template_path
from .template_renderer import render_docxtpl_template


def sanitize_filename(name: str) -> str:
    forbidden = '<>:"/\\|?*'
    cleaned = "".join("_" if ch in forbidden else ch for ch in name).strip()
    return cleaned or "document"


def build_docx_bytes(payload: DocumentPreviewPayload) -> bytes:
    template_path = get_template_path(payload.templateKey)

    if template_path and template_path.exists():
        return render_docxtpl_template(template_path, payload)

    return build_fallback_docx_bytes(payload)


def build_pdf_bytes(payload: DocumentPreviewPayload) -> bytes:
    docx_bytes = build_docx_bytes(payload)
    filename_stem = (
        sanitize_filename(payload.outputFilename)
        .replace(".docx", "")
        .replace(".pdf", "")
    )
    return convert_docx_bytes_to_pdf_bytes(docx_bytes, filename_stem)


def build_document_bytes(
    payload: DocumentPreviewPayload,
    output_format: OutputFormat,
) -> tuple[bytes, str, str]:
    safe_name = (
        sanitize_filename(payload.outputFilename)
        .replace(".docx", "")
        .replace(".pdf", "")
    )

    if output_format == "pdf":
        pdf_bytes = build_pdf_bytes(payload)
        return pdf_bytes, "application/pdf", f"{safe_name}.pdf"

    docx_bytes = build_docx_bytes(payload)
    return (
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{safe_name}.docx",
    )


def build_package_zip_bytes(
    payloads: list[DocumentPreviewPayload],
    output_formats: Sequence[OutputFormat] | None = None,
) -> bytes:
    """
    ZIP package builder.

    Default behavior:
    - if output_formats is None, package both DOCX and PDF for each payload.
    """
    effective_formats: Sequence[OutputFormat] = output_formats or ("docx", "pdf")

    zip_buffer = BytesIO()

    with ZipFile(zip_buffer, mode="w", compression=ZIP_DEFLATED) as zip_file:
        for payload in payloads:
            for output_format in effective_formats:
                if output_format == "zip":
                    continue

                file_bytes, _media_type, filename = build_document_bytes(
                    payload, output_format
                )
                zip_file.writestr(filename, file_bytes)

    return zip_buffer.getvalue()


def build_fallback_docx_bytes(payload: DocumentPreviewPayload) -> bytes:
    doc = Document()

    doc.add_heading(payload.documentName, level=0)

    meta = doc.add_table(rows=0, cols=2)
    _add_row(meta, "Template Key", payload.templateKey)
    _add_row(meta, "Deal ID", payload.dealId)
    _add_row(meta, "Business Name", payload.businessName)
    _add_row(meta, "Generated At", payload.generatedAt)
    _add_row(meta, "Output Filename", payload.outputFilename)

    doc.add_paragraph("")
    doc.add_heading("Template Payload", level=1)

    _render_mapping(doc, payload.data, level=2)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_mapping(doc: Document, value: dict[str, Any], level: int = 2) -> None:
    for key, item in value.items():
        heading = doc.add_paragraph()
        heading.style = f"Heading {min(level, 9)}"
        heading.add_run(_prettify_key(key))

        _render_value(doc, item, level + 1)


def _render_value(doc: Document, value: Any, level: int) -> None:
    if isinstance(value, dict):
        _render_mapping(doc, value, level)
        return

    if isinstance(value, list):
        _render_list(doc, value, level)
        return

    paragraph = doc.add_paragraph()
    paragraph.add_run(_stringify(value))


def _render_list(doc: Document, values: Iterable[Any], level: int) -> None:
    values = list(values)

    if not values:
        doc.add_paragraph("(empty)")
        return

    for item in values:
        if isinstance(item, dict):
            sub_doc = doc.add_paragraph(style="List Bullet")
            sub_doc.add_run("Item")
            _render_mapping(doc, item, level + 1)
        elif isinstance(item, list):
            doc.add_paragraph("Nested List", style="List Bullet")
            _render_list(doc, item, level + 1)
        else:
            doc.add_paragraph(_stringify(item), style="List Bullet")


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    return str(value)


def _prettify_key(key: str) -> str:
    return key.replace("_", " ").strip().title()


def _add_row(table, left: str, right: str) -> None:
    row = table.add_row().cells
    row[0].text = left
    row[1].text = right