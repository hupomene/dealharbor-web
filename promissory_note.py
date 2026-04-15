from docx import Document

doc = Document()

doc.add_heading("PROMISSORY NOTE", 0)

doc.add_paragraph("Principal Amount: {{ principal_amount }}")
doc.add_paragraph("Date: {{ issue_date }}")
doc.add_paragraph("Borrower: {{ borrower_name }}")
doc.add_paragraph("Lender: {{ lender_name }}")
doc.add_paragraph("Business: {{ business_name }}")
doc.add_paragraph("State: {{ state }}")

doc.add_heading("1. Promise to Pay", level=1)
doc.add_paragraph(
    "For value received, the Borrower promises to pay to the Lender "
    "the principal sum of {{ principal_amount }} under the terms of this Note."
)

doc.add_heading("2. Payment Terms", level=1)
doc.add_paragraph(
    "Repayment terms, interest rate, and amortization schedule to be defined."
)

doc.add_heading("3. Default", level=1)
doc.add_paragraph(
    "Failure to make payments constitutes default."
)

doc.add_heading("4. Governing Law", level=1)
doc.add_paragraph(
    "This Note shall be governed by the laws of {{ state }}."
)

doc.add_heading("Signatures", level=1)
doc.add_paragraph("Borrower: ________________________")
doc.add_paragraph("Lender: ________________________")

doc.save("promissory_note.docx")